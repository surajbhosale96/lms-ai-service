import io
import requests
import mammoth
from docx import Document


def extract_docx(url: str) -> str:
    response = requests.get(url, timeout=60)
    response.raise_for_status()

    content = io.BytesIO(response.content)

    # Try .docx first (python-docx)
    try:
        doc = Document(content)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        if paragraphs:
            return "\n\n".join(paragraphs)
    except Exception:
        pass

    # Fallback: mammoth handles both .doc and .docx
    content.seek(0)
    result = mammoth.extract_raw_text(content)
    return result.value.strip()
